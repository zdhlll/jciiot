"""Generate SOP Markdown knowledge files from the original DOCX manuals.

This script is intentionally standalone. It does not import or read the
existing `knowledge/sop*.md` files. The Markdown shape is reproduced from a
local template so generation can be audited as code.
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


FORBIDDEN_REFERENCE_SOPS = {
    "knowledge/sop1.md",
    "knowledge/sop2.md",
    "knowledge/sop3.md",
    "knowledge/sop4.md",
    "knowledge/sop5.md",
    "knowledge/sop_main.md",
}


@dataclass(frozen=True)
class CaseSpec:
    case_no: int
    level: str
    level_index: int
    scene_short: str
    scene_prefix: str
    max_score: int
    doc_name: str
    output_name: str
    title_suffix: str


CASES: tuple[CaseSpec, ...] = (
    CaseSpec(
        1, "L1", 0, "factory_sorting_1", "factory_sorting_1_3fo3erfhisem",
        10, "JCIIOT 2026 case 1 SOP.docx", "sop1.md",
        "Single-line Transport",
    ),
    CaseSpec(
        3, "L2", 1, "factory_sorting_3", "factory_sorting_3_3fo3errph7x9",
        15, "JCIIOT 2026 case 3 SOP.docx", "sop2.md",
        "Cross-line Transport + Obstacle Avoidance",
    ),
    CaseSpec(
        5, "L3", 2, "factory_sorting_5", "factory_sorting_5_3fo3ertpxeut",
        20, "JCIIOT 2026 case 5 SOP.docx", "sop3.md",
        "Cross-line Transport + Obstacle + Interference",
    ),
    CaseSpec(
        7, "L4", 3, "factory_sorting_7", "factory_sorting_7_3fo3erfky9rn",
        25, "JCIIOT 2026 case 7 SOP.docx", "sop4.md",
        "Large Object Cross-line Transport",
    ),
    CaseSpec(
        9, "L5", 4, "factory_sorting_9", "factory_sorting_9_3fo3ert2c5fp",
        30, "JCIIOT 2026 case 9 SOP.docx", "sop5.md",
        "Extreme Distance Transport",
    ),
)


GRASP_POSES: dict[str, tuple[float, float, float]] = {
    "input_1": (5.03, -3.84, -3.14),
    "input_2": (8.56, -3.92, -3.14),
    "input_3": (12.38, -3.76, -3.14),
    "input_4": (15.80, -3.77, -3.14),
    "input_5": (8.00, 4.60, -3.139),
    "input_6": (6.00, 4.80, -3.139),
}


@dataclass
class DocxPayload:
    file: Path
    paragraphs: list[str]
    tables: list[str]
    images: list[tuple[str, bytes]]

    @property
    def text(self) -> str:
        blocks = list(self.paragraphs)
        if self.tables:
            blocks.append("\n".join(self.tables))
        return "\n\n".join(blocks)


@dataclass
class SopFacts:
    level: str
    scene: str
    max_score: int
    task_summary: str = ""
    object_description: str = ""
    pick_station_label: str = ""
    place_station_label: str = ""
    source: str = ""
    target: str = ""
    object_names: list[str] = field(default_factory=list)
    inventory: dict[str, list[str]] = field(default_factory=dict)
    constraints: list[str] = field(default_factory=list)
    image_findings: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    llm_used: bool = False
    vlm_used: bool = False


def relpath(path: Path, root: Path) -> str:
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return path.as_posix()


def ensure_no_reference_sop_inputs(project_root: Path, paths: list[Path]) -> None:
    forbidden = {project_root / item for item in FORBIDDEN_REFERENCE_SOPS}
    for path in paths:
        resolved = path.resolve()
        if any(resolved == f.resolve() for f in forbidden):
            raise RuntimeError(f"Forbidden SOP reference file used as input: {path}")


def read_text_if_exists(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def load_docx(path: Path) -> DocxPayload:
    if path.name.startswith("~$"):
        raise ValueError(f"Refusing temporary Word lock file: {path}")
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install project requirements first: "
            "pip install -r requirements.txt"
        ) from exc

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[str] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            tables.append(f"[table {table_index}]\n" + "\n".join(rows))

    images: list[tuple[str, bytes]] = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            name = rel.target_ref.split("/")[-1] if rel.target_ref else "image.png"
            images.append((name, rel.target_part.blob))
    return DocxPayload(file=path, paragraphs=paragraphs, tables=tables, images=images)


def guess_image_mime(image_bytes: bytes) -> str:
    if image_bytes[:4] == b"\x89PNG":
        return "image/png"
    if image_bytes[:2] == b"\xff\xd8":
        return "image/jpeg"
    if image_bytes[:4] == b"RIFF" and image_bytes[8:12] == b"WEBP":
        return "image/webp"
    return "image/png"


def post_json(url: str, payload: dict[str, Any], headers: dict[str, str], timeout: float) -> dict[str, Any]:
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers=headers,
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return json.loads(resp.read().decode("utf-8"))


def call_openai_chat(
    messages: list[dict[str, Any]],
    *,
    base_url: str,
    api_key: str,
    model: str,
    temperature: float,
    max_tokens: int,
    timeout: float,
    json_mode: bool = False,
) -> str:
    endpoint = f"{base_url.rstrip('/')}/chat/completions"
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    payload: dict[str, Any] = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    data = post_json(endpoint, payload, headers, timeout)
    return data["choices"][0]["message"]["content"]


def call_ollama_chat(
    messages: list[dict[str, Any]],
    *,
    base_url: str,
    model: str,
    temperature: float,
    timeout: float,
    json_mode: bool = False,
) -> str:
    payload: dict[str, Any] = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {"temperature": temperature},
    }
    if json_mode:
        payload["format"] = "json"
    data = post_json(
        f"{base_url.rstrip('/')}/api/chat",
        payload,
        {"Content-Type": "application/json"},
        timeout,
    )
    return data["message"]["content"]


def model_provider() -> str:
    explicit = os.getenv("SOP_LLM_PROVIDER", "").strip().lower()
    if explicit:
        return explicit
    if os.getenv("OPENAI_API_KEY", ""):
        return "openai"
    return "ollama"


def call_text_model(prompt: str, *, require_json: bool = True, timeout: float = 120.0) -> str:
    provider = model_provider()
    messages = [
        {
            "role": "system",
            "content": (
                "You extract factory robot SOP facts from DOCX text. "
                "Return only the requested JSON object."
            ),
        },
        {"role": "user", "content": prompt},
    ]
    if provider == "openai":
        return call_openai_chat(
            messages,
            base_url=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
            api_key=os.getenv("OPENAI_API_KEY", ""),
            model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            temperature=float(os.getenv("SOP_LLM_TEMPERATURE", "0.0")),
            max_tokens=int(os.getenv("SOP_LLM_MAX_TOKENS", "4096")),
            timeout=timeout,
            json_mode=require_json,
        )
    return call_ollama_chat(
        messages,
        base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
        model=os.getenv("OLLAMA_MODEL", "qwen3.6:27b-mtp-q4_K_M"),
        temperature=float(os.getenv("SOP_LLM_TEMPERATURE", "0.0")),
        timeout=timeout,
        json_mode=require_json,
    )


def call_vision_model(prompt: str, image_bytes: bytes, *, timeout: float = 90.0) -> str:
    base_url = os.getenv("VLM_BASE_URL") or os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    model = os.getenv("VLM_MODEL", "qwen3-vl:8b")
    api_key = os.getenv("VLM_API_KEY") or os.getenv("OPENAI_API_KEY", "")
    provider = os.getenv("VLM_PROVIDER", "").strip().lower()
    if not provider:
        provider = "openai" if api_key and ("openai" in base_url.lower() or "/v1" in base_url.lower()) else "ollama"

    if provider == "openai":
        b64 = base64.b64encode(image_bytes).decode("ascii")
        mime = guess_image_mime(image_bytes)
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                ],
            }
        ]
        return call_openai_chat(
            messages,
            base_url=base_url,
            api_key=api_key,
            model=model,
            temperature=0.0,
            max_tokens=800,
            timeout=timeout,
            json_mode=False,
        )

    b64 = base64.b64encode(image_bytes).decode("ascii")
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt, "images": [b64]}],
        "stream": False,
    }
    data = post_json(
        f"{base_url.rstrip('/')}/api/chat",
        payload,
        {"Content-Type": "application/json"},
        timeout,
    )
    return data["message"]["content"]


def describe_docx_images(payload: DocxPayload, *, enabled: bool) -> tuple[list[str], bool]:
    if not enabled or not payload.images:
        return [], False
    descriptions: list[str] = []
    for name, image_bytes in payload.images:
        try:
            desc = call_vision_model(
                "Describe this SOP figure for a factory material handling robot. "
                "Focus on pick/place station labels, object appearance, route, "
                "obstacles, and anything relevant to the task.",
                image_bytes,
            )
            descriptions.append(f"{name}: {desc.strip()}")
        except Exception as exc:
            descriptions.append(f"{name}: VLM unavailable ({exc})")
    return descriptions, any("VLM unavailable" not in d for d in descriptions)


def extract_json_object(raw: str) -> dict[str, Any]:
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, flags=re.S)
    if match:
        return json.loads(match.group(1))
    start = raw.find("{")
    end = raw.rfind("}")
    if start >= 0 and end > start:
        return json.loads(raw[start : end + 1])
    raise ValueError(f"No JSON object found in model output: {raw[:200]}")


def station_label_to_port(label: str, role: str) -> str:
    match = re.search(r"(\d+)", label or "")
    if not match:
        return ""
    station_no = int(match.group(1))
    if not 1 <= station_no <= 6:
        return ""
    prefix = "input" if role == "pick" else "output"
    # In the published factory layout, external station numbering is opposite
    # the simulator port index: Station 1 maps to *_6, Station 6 maps to *_1.
    return f"{prefix}_{7 - station_no}"


def port_to_station_label(port: str, role: str) -> str:
    match = re.search(r"_(\d+)$", port or "")
    if not match:
        return ""
    station_no = 7 - int(match.group(1))
    prefix = "Pick Station" if role == "pick" else "Place Station"
    return f"{prefix} {station_no}"


def compact_text(text: str, limit: int) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= limit:
        return text
    head = text[: limit // 2]
    tail = text[-limit // 2 :]
    return f"{head}\n\n[... middle omitted for token budget ...]\n\n{tail}"


def build_llm_prompt(spec: CaseSpec, docx: DocxPayload, erratum: str, image_descriptions: list[str]) -> str:
    images = "\n".join(f"- {d}" for d in image_descriptions) or "No image descriptions available."
    return f"""
Extract SOP facts for {spec.level} / case {spec.case_no}.

Important constraints:
- Use only the DOCX text, DOCX image descriptions, and erratum below.
- Do not use any prewritten Markdown SOP knowledge file.
- If the erratum corrects a station reference, apply the correction.
- Preserve exact object identifiers with underscores if the document or image contains them.
- Return JSON only.

Expected JSON schema:
{{
  "task_summary": "short English task sentence",
  "object_description": "object type, color, and quantity",
  "pick_station_label": "Pick Station N or Placement Point N",
  "place_station_label": "Place Station N",
  "object_names": ["exact_object_name_if_known"],
  "inventory": {{"input_1": ["object_a"], "input_2": []}},
  "constraints": ["short constraint"],
  "image_findings": ["short image finding"],
  "confidence": 0.0
}}

Erratum:
{compact_text(erratum, 4000)}

DOCX text:
{compact_text(docx.text, 14000)}

DOCX image descriptions:
{compact_text(images, 5000)}
""".strip()


def regex_fallback_facts(spec: CaseSpec, text: str) -> SopFacts:
    facts = SopFacts(level=spec.level, scene=spec.scene_short, max_score=spec.max_score)
    normalized = " ".join(text.split())

    pick_patterns = [
        r"starting point\s+\"([^\"]*Pick Station\s+\d+)\"",
        r"Pick Station\s+is\s+(Pick Station\s+\d+|Placement Point\s+\d+)",
        r"Starting Location:\s*(Pick Station\s+\d+|Placement Point\s+\d+)",
        r"from\s+(Pick Station\s+\d+|Placement Point\s+\d+)",
    ]
    place_patterns = [
        r"destination\s+\"([^\"]*Place Station\s+\d+)\"",
        r"Place Station\s+is\s+(Place Station\s+\d+)",
        r"Target Location:\s*(Place Station\s+\d+)",
        r"to\s+(Place Station\s+\d+)",
    ]
    for pattern in pick_patterns:
        match = re.search(pattern, normalized, flags=re.I)
        if match:
            facts.pick_station_label = cleanup_station_label(match.group(1), "pick")
            break
    for pattern in place_patterns:
        match = re.search(pattern, normalized, flags=re.I)
        if match:
            facts.place_station_label = cleanup_station_label(match.group(1), "place")
            break

    object_patterns = [
        r"object(?: to be handled)? is (.*?)(?:\.|, and| The Pick|$)",
        r"transport (?:a|an|the|three) (.*?)(?: from|\.|$)",
        r"Move the (.*?) from",
        r"Material Name:\s*([^\.]+?)(?: Starting Location|$)",
    ]
    for pattern in object_patterns:
        match = re.search(pattern, normalized, flags=re.I)
        if match:
            facts.object_description = match.group(1).strip(" .")
            break

    if facts.pick_station_label:
        facts.source = station_label_to_port(facts.pick_station_label, "pick")
    if facts.place_station_label:
        facts.target = station_label_to_port(facts.place_station_label, "place")
    facts.task_summary = build_task_sentence(facts)
    facts.constraints = [
        "Follow the SOP sequence: move to pick station, pick material, move to place station, place material.",
        "Confirm grasp before transport.",
        "Avoid obstacles and collisions during navigation.",
    ]
    return facts


def cleanup_station_label(label: str, role: str) -> str:
    label = re.sub(r"\s+", " ", label or "").strip().strip(".")
    match = re.search(r"(\d+)", label)
    if not match:
        return label
    number = match.group(1)
    if role == "pick":
        return f"Pick Station {number}"
    return f"Place Station {number}"


def facts_from_model(spec: CaseSpec, raw: dict[str, Any]) -> SopFacts:
    facts = SopFacts(level=spec.level, scene=spec.scene_short, max_score=spec.max_score, llm_used=True)
    facts.task_summary = str(raw.get("task_summary") or "").strip()
    facts.object_description = str(raw.get("object_description") or "").strip()
    facts.pick_station_label = cleanup_station_label(str(raw.get("pick_station_label") or ""), "pick")
    facts.place_station_label = cleanup_station_label(str(raw.get("place_station_label") or ""), "place")

    object_names = raw.get("object_names") or []
    if isinstance(object_names, str):
        object_names = [object_names]
    facts.object_names = [str(item).strip() for item in object_names if str(item).strip()]

    inventory = raw.get("inventory") or {}
    if isinstance(inventory, dict):
        for key, value in inventory.items():
            if isinstance(value, str):
                items = [v.strip() for v in value.split(",") if v.strip()]
            elif isinstance(value, list):
                items = [str(v).strip() for v in value if str(v).strip()]
            else:
                items = []
            if items:
                facts.inventory[str(key).strip()] = items

    for field_name, target in (("constraints", facts.constraints), ("image_findings", facts.image_findings)):
        value = raw.get(field_name) or []
        if isinstance(value, str):
            value = [value]
        if isinstance(value, list):
            target.extend(str(item).strip() for item in value if str(item).strip())

    facts.source = station_label_to_port(facts.pick_station_label, "pick")
    facts.target = station_label_to_port(facts.place_station_label, "place")
    if not facts.task_summary:
        facts.task_summary = build_task_sentence(facts)
    return facts


def load_task_config(project_root: Path) -> dict[str, Any]:
    path = project_root / "knowledge" / "task_config.json"
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def task_config_for_spec(task_config: dict[str, Any], spec: CaseSpec) -> dict[str, Any]:
    tasks = task_config.get("tasks") or []
    for task in tasks:
        if task.get("level") == spec.level:
            return dict(task)
    if spec.level_index < len(tasks):
        return dict(tasks[spec.level_index])
    return {}


def load_semantic_map(project_root: Path, spec: CaseSpec) -> dict[str, Any]:
    path = (
        project_root / "robosuite" / "robosuite" / "environments" / "factory_sorting"
        / "generated_maps" / f"{spec.scene_prefix}_scene_regenerated_semantic_map.json"
    )
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def station_center(semantic_map: dict[str, Any], port: str) -> tuple[float, float] | None:
    for block_name in ("input_ports", "output_ports"):
        block = semantic_map.get(block_name) or {}
        if port in block and isinstance(block[port], dict):
            center = block[port].get("center")
            if isinstance(center, list) and len(center) >= 2:
                return float(center[0]), float(center[1])
    for obj in semantic_map.get("objects") or []:
        if isinstance(obj, dict) and obj.get("name") == port:
            center = obj.get("center")
            if isinstance(center, list) and len(center) >= 2:
                return float(center[0]), float(center[1])
    return None


def apply_validation_and_fallbacks(
    facts: SopFacts,
    spec: CaseSpec,
    task_cfg: dict[str, Any],
    semantic_map: dict[str, Any],
    *,
    allow_task_config_fill: bool,
) -> SopFacts:
    cfg_source = str(task_cfg.get("source") or "")
    cfg_target = str(task_cfg.get("target") or "")
    cfg_object = str(task_cfg.get("object") or "")

    if not facts.source and facts.pick_station_label:
        facts.source = station_label_to_port(facts.pick_station_label, "pick")
    if not facts.target and facts.place_station_label:
        facts.target = station_label_to_port(facts.place_station_label, "place")
    if not facts.pick_station_label and cfg_source and allow_task_config_fill:
        facts.pick_station_label = port_to_station_label(cfg_source, "pick")
        facts.source = cfg_source
        facts.warnings.append("Filled missing pick station from task_config validation data.")
    if not facts.place_station_label and cfg_target and allow_task_config_fill:
        facts.place_station_label = port_to_station_label(cfg_target, "place")
        facts.target = cfg_target
        facts.warnings.append("Filled missing place station from task_config validation data.")

    if cfg_source and facts.source and facts.source != cfg_source:
        facts.warnings.append(f"Source mismatch: docx/model derived {facts.source}, task_config has {cfg_source}.")
        if allow_task_config_fill:
            facts.source = cfg_source
            facts.pick_station_label = port_to_station_label(cfg_source, "pick")
            facts.warnings.append("Corrected source from task_config validation data.")
    if cfg_target and facts.target and facts.target != cfg_target:
        facts.warnings.append(f"Target mismatch: docx/model derived {facts.target}, task_config has {cfg_target}.")
        if allow_task_config_fill:
            facts.target = cfg_target
            facts.place_station_label = port_to_station_label(cfg_target, "place")
            facts.warnings.append("Corrected target from task_config validation data.")

    has_exact_object = any("_" in name for name in facts.object_names)
    if not has_exact_object and cfg_object and allow_task_config_fill:
        facts.object_names = [cfg_object]
        facts.warnings.append("Filled missing exact object_name from task_config validation data.")

    if facts.source and facts.object_names:
        inv = facts.inventory.setdefault(facts.source, [])
        for obj in facts.object_names:
            if obj not in inv:
                inv.append(obj)

    if not facts.task_summary:
        facts.task_summary = build_task_sentence(facts)
    if not facts.constraints:
        facts.constraints = [
            "Follow the SOP order exactly.",
            "pick_up requires both target and exact object_name.",
            "Avoid obstacles and collisions.",
        ]
    return facts


def build_task_sentence(facts: SopFacts) -> str:
    obj = facts.object_description or ("; ".join(facts.object_names) if facts.object_names else "the target material")
    pick = facts.pick_station_label or facts.source or "the pick station"
    place = facts.place_station_label or facts.target or "the place station"
    return f"Transport {obj} from {pick} to {place}"


def fmt_point(point: tuple[float, float] | None) -> str:
    if point is None:
        return ""
    return f"({point[0]:.2f}, {point[1]:.2f})"


def render_sop_markdown(facts: SopFacts, spec: CaseSpec, semantic_map: dict[str, Any], source_doc: Path) -> str:
    source_center = station_center(semantic_map, facts.source)
    target_center = station_center(semantic_map, facts.target)
    object_lines: list[str]
    if len(facts.object_names) > 1:
        object_lines = ["- Target objects:"] + [f"  - {obj}" for obj in facts.object_names]
    elif facts.object_names:
        object_lines = [f"- Target object: {facts.object_names[0]}"]
    else:
        object_lines = ["- Target object: UNKNOWN"]

    lines: list[str] = [
        "<!-- AI-GENERATED FROM DOCX - DO NOT MODIFY MANUALLY -->",
        "",
        f"# {spec.level} Task - {spec.title_suffix}",
        "",
        f"Level: {spec.level} (max {spec.max_score} points)",
        f"Scene: {spec.scene_short}",
        "",
        "## Task",
        "",
        facts.task_summary.rstrip("."),
        "",
        "## Station Mapping",
        "",
        f"- {facts.pick_station_label or 'Pick Station UNKNOWN'} = {facts.source or 'UNKNOWN'}"
        + (f", center {fmt_point(source_center)}" if source_center else ""),
        f"- {facts.place_station_label or 'Place Station UNKNOWN'} = {facts.target or 'UNKNOWN'}"
        + (f", center {fmt_point(target_center)}" if target_center else ""),
        "- Robot start: (13.5, 0.0)",
    ]
    lines.extend(object_lines)

    if facts.source in GRASP_POSES:
        x, y, yaw = GRASP_POSES[facts.source]
        lines.extend([
            "",
            "## Grasp Pose (BC Policy)",
            "",
            f"- Robot stop point at {facts.source}: ({x:.2f}, {y:.2f}, 0.0), yaw={yaw:.3f}",
        ])

    if facts.inventory:
        lines.extend([
            "",
            f"## Object Inventory ({spec.level} Scene)",
            "",
            "Every input port and its assigned graspable object:",
            "",
        ])
        for input_id in [f"input_{i}" for i in range(1, 7)]:
            objects = facts.inventory.get(input_id) or []
            if objects:
                lines.append(f"- {input_id}: {', '.join(objects)}")

    lines.extend([
        "",
        "CRITICAL: When calling pick_up, you MUST provide the exact object_name from the inventory above.",
    ])

    if facts.constraints:
        lines.extend(["", "## SOP Constraints", ""])
        for item in facts.constraints[:8]:
            lines.append(f"- {item.rstrip('.')}.")

    lines.extend([
        "",
        "## Generation Evidence",
        "",
        f"- Source DOCX: `{source_doc.as_posix()}`",
        "- Runtime policy: existing `knowledge/sop*.md` files were not read.",
    ])
    if facts.warnings:
        lines.extend(["", "## Generation Warnings", ""])
        for warning in facts.warnings:
            lines.append(f"- {warning}")

    return "\n".join(lines).rstrip() + "\n"


def render_sop_main(rows: list[tuple[CaseSpec, SopFacts, dict[str, Any]]]) -> str:
    lines = [
        "<!-- AI-GENERATED FROM DOCX - DO NOT MODIFY MANUALLY -->",
        "",
        "# Standard Operating Procedure (SOP)",
        "",
        "Task ID: MT-MOBILE-001",
        "Version: v2.0",
        "",
        "## Standard Transport Workflow",
        "",
        "1. Navigate to Pick Station",
        "2. Pick material: move arm above object -> close gripper -> lift 150mm -> confirm grasp",
        "3. Navigate to Place Station with object held",
        "4. Place material: lower to table height -> open gripper -> confirm deviation < 10mm",
        "5. Return or repeat",
        "",
        "## Task Coordinate Reference",
        "",
        "All five levels, robot start at (13.5, 0.0):",
        "",
        "| Level | Scene | Pick Station | Pick Coords | Object Name | Place Station | Place Coords |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for spec, facts, semantic in rows:
        source_center = station_center(semantic, facts.source)
        target_center = station_center(semantic, facts.target)
        objects = "; ".join(facts.object_names) if facts.object_names else "UNKNOWN"
        lines.append(
            f"| {spec.level} | {spec.scene_short} | {facts.pick_station_label or 'UNKNOWN'} | "
            f"{facts.source or 'UNKNOWN'} {fmt_point(source_center)} | {objects} | "
            f"{facts.place_station_label or 'UNKNOWN'} | {facts.target or 'UNKNOWN'} {fmt_point(target_center)} |"
        )
    lines.extend([
        "",
        "## CRITICAL pick_up Rules",
        "",
        "- pick_up requires BOTH `target` (station name like input_6) AND `object_name` (exact object name from the table above)",
        "- Never guess object names - always use the exact name from the per-level SOP Object Inventory",
        "",
        "## BC Policy Grasp Poses",
        "",
        "| Input Station | Grasp Pose (x, y, yaw) |",
        "| --- | --- |",
    ])
    for input_id, pose in GRASP_POSES.items():
        lines.append(f"| {input_id} | ({pose[0]:.2f}, {pose[1]:.2f}, {pose[2]:.3f}) |")
    return "\n".join(lines).rstrip() + "\n"


def write_index(out_dir: Path, files: list[Path]) -> None:
    docs: dict[str, Any] = {}
    for file in sorted(files):
        content = file.read_text(encoding="utf-8")
        title = next((line.lstrip("# ").strip() for line in content.splitlines() if line.startswith("# ")), file.stem)
        tags = re.findall(r"\bL[1-5]\b|input_\d|output_\d", content)
        docs[file.name] = {
            "title": title,
            "category": "sop",
            "tags": list(dict.fromkeys(tags))[:8],
            "added_at": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime()),
        }
    payload = {
        "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime()),
        "document_count": len(docs),
        "documents": docs,
    }
    (out_dir / "_index.json").write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def generate_one(
    project_root: Path,
    spec: CaseSpec,
    *,
    erratum: str,
    task_config: dict[str, Any],
    use_vision: bool,
    require_llm: bool,
    allow_task_config_fill: bool,
) -> tuple[SopFacts, dict[str, Any], str]:
    docx_path = project_root / "sop+prompt" / spec.doc_name
    ensure_no_reference_sop_inputs(project_root, [docx_path])
    docx = load_docx(docx_path)
    image_descriptions, vlm_used = describe_docx_images(docx, enabled=use_vision)

    facts: SopFacts
    llm_error = ""
    try:
        prompt = build_llm_prompt(spec, docx, erratum, image_descriptions)
        raw = call_text_model(prompt, require_json=True)
        facts = facts_from_model(spec, extract_json_object(raw))
        facts.vlm_used = vlm_used
    except Exception as exc:
        llm_error = str(exc)
        if require_llm:
            raise
        facts = regex_fallback_facts(spec, docx.text)
        facts.vlm_used = vlm_used
        facts.warnings.append(f"LLM extraction unavailable; deterministic fallback used: {llm_error}")

    semantic_map = load_semantic_map(project_root, spec)
    task_cfg = task_config_for_spec(task_config, spec)
    facts = apply_validation_and_fallbacks(
        facts,
        spec,
        task_cfg,
        semantic_map,
        allow_task_config_fill=allow_task_config_fill,
    )
    md = render_sop_markdown(facts, spec, semantic_map, Path("sop+prompt") / spec.doc_name)
    return facts, semantic_map, md


def generate_all(args: argparse.Namespace) -> int:
    project_root = Path(args.project_root).resolve()
    out_dir = Path(args.out_dir)
    if not out_dir.is_absolute():
        out_dir = project_root / out_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    erratum_path = project_root.parent / "ERRATUM.md"
    erratum = read_text_if_exists(erratum_path)
    task_config = load_task_config(project_root)

    generated_files: list[Path] = []
    rows: list[tuple[CaseSpec, SopFacts, dict[str, Any]]] = []
    report: dict[str, Any] = {
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime()),
        "project_root": str(project_root),
        "out_dir": str(out_dir),
        "forbidden_reference_sops": sorted(FORBIDDEN_REFERENCE_SOPS),
        "cases": [],
    }

    for spec in CASES:
        facts, semantic_map, md = generate_one(
            project_root,
            spec,
            erratum=erratum,
            task_config=task_config,
            use_vision=not args.no_vision,
            require_llm=args.require_llm,
            allow_task_config_fill=not args.strict_docx_only,
        )
        out_file = out_dir / spec.output_name
        out_file.write_text(md, encoding="utf-8")
        generated_files.append(out_file)
        rows.append((spec, facts, semantic_map))
        report["cases"].append({
            "level": spec.level,
            "docx": spec.doc_name,
            "output": str(out_file),
            "llm_used": facts.llm_used,
            "vlm_used": facts.vlm_used,
            "source": facts.source,
            "target": facts.target,
            "object_names": facts.object_names,
            "warnings": facts.warnings,
        })

    sop_main = out_dir / "sop_main.md"
    sop_main.write_text(render_sop_main(rows), encoding="utf-8")
    generated_files.append(sop_main)
    write_index(out_dir, generated_files)
    (out_dir / "generation_report.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    print(f"Generated {len(generated_files)} SOP files in {out_dir}")
    print(f"Report: {out_dir / 'generation_report.json'}")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate SOP Markdown files from DOCX manuals.")
    parser.add_argument("--project-root", default=".", help="JCIIOT project root.")
    parser.add_argument("--out-dir", default="sop_autogen_tool/generated_knowledge", help="Output directory.")
    parser.add_argument("--no-vision", action="store_true", help="Skip VLM image descriptions.")
    parser.add_argument("--require-llm", action="store_true", help="Fail instead of using deterministic fallback.")
    parser.add_argument(
        "--strict-docx-only",
        action="store_true",
        help="Do not fill missing source/target/object fields from task_config validation data.",
    )
    args = parser.parse_args(argv)
    try:
        return generate_all(args)
    except (RuntimeError, ValueError, urllib.error.URLError, TimeoutError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
